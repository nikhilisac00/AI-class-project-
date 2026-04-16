"""
Memo Export — PDF and DOCX generation from the IC Due Diligence Memo.

Usage:
    from tools.memo_export import to_pdf, to_docx
    pdf_bytes  = to_pdf(memo_markdown, firm_name="AQR Capital")
    docx_bytes = to_docx(memo_markdown, firm_name="AQR Capital")
"""

from __future__ import annotations
import io
import re


# ── Shared markdown parser ────────────────────────────────────────────────────

def _parse_table(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into a 2-D list of strings."""
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-:| ]+\|\s*$", line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells:
            rows.append(cells)
    return rows


def _strip_md(text: str) -> str:
    """Remove inline markdown (bold, italic, links) for plain-text contexts."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text.strip()


def _inline_bold(text: str) -> str:
    """Convert **bold** to ReportLab <b> tags."""
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)


# ── PDF via ReportLab ─────────────────────────────────────────────────────────

def to_pdf(memo_markdown: str, summary: str = "", firm_name: str = "") -> bytes:
    """Convert a markdown IC memo to a professionally styled PDF."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, KeepTogether,
    )

    # ── Palette ───────────────────────────────────────────────────────────────
    NAVY       = colors.HexColor("#0a2744")
    BLUE       = colors.HexColor("#1a5fa8")
    LIGHT_BLUE = colors.HexColor("#dce8f5")
    DARK_GRAY  = colors.HexColor("#1e1e1e")
    MID_GRAY   = colors.HexColor("#555555")
    RULE_GRAY  = colors.HexColor("#cbd5e0")
    TH_BG      = colors.HexColor("#1a3a5c")
    ALT_ROW    = colors.HexColor("#f3f7fb")
    WHITE      = colors.white

    SEV_COLORS = {
        "HIGH":     colors.HexColor("#dc3545"),
        "MEDIUM":   colors.HexColor("#fd7e14"),
        "LOW":      colors.HexColor("#e6a817"),
        "CRITICAL": colors.HexColor("#9b1c2e"),
        "CLEAN":    colors.HexColor("#28a745"),
        "PASS":     colors.HexColor("#28a745"),
        "WARN":     colors.HexColor("#fd7e14"),
        "INFO":     colors.HexColor("#0d6efd"),
    }

    buf = io.BytesIO()

    # ── Header / footer callback ──────────────────────────────────────────────
    _fname = firm_name.upper() if firm_name else "INVESTMENT MANAGER"

    def _on_page(canvas, doc):
        canvas.saveState()
        w, h = LETTER
        # Top bar
        canvas.setFillColor(NAVY)
        canvas.rect(0, h - 0.55 * inch, w, 0.55 * inch, fill=1, stroke=0)
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColor(WHITE)
        canvas.drawString(0.65 * inch, h - 0.35 * inch,
                          "INVESTMENT COMMITTEE  ·  DUE DILIGENCE MEMO  ·  CONFIDENTIAL DRAFT")
        canvas.drawRightString(w - 0.5 * inch, h - 0.35 * inch, _fname)
        # Bottom rule + footer text
        canvas.setStrokeColor(RULE_GRAY)
        canvas.setLineWidth(0.4)
        canvas.line(0.5 * inch, 0.6 * inch, w - 0.5 * inch, 0.6 * inch)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(MID_GRAY)
        canvas.drawString(0.5 * inch, 0.38 * inch,
                          "CONFIDENTIAL — DRAFT FOR IC REVIEW ONLY. Not for distribution.")
        canvas.drawRightString(w - 0.5 * inch, 0.38 * inch, f"Page {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.85 * inch,
    )

    # ── Styles ────────────────────────────────────────────────────────────────
    H1 = ParagraphStyle("H1", fontName="Helvetica-Bold", fontSize=18,
                         textColor=NAVY, spaceAfter=4, spaceBefore=8, leading=22,
                         alignment=TA_CENTER)
    H1_SUB = ParagraphStyle("H1Sub", fontName="Helvetica", fontSize=9,
                              textColor=MID_GRAY, spaceAfter=6, leading=13,
                              alignment=TA_CENTER)
    H2 = ParagraphStyle("H2", fontName="Helvetica-Bold", fontSize=11.5,
                         textColor=WHITE, spaceAfter=6, spaceBefore=14, leading=15,
                         leftIndent=6, backColor=TH_BG, borderPad=(4, 4, 4, 6))
    H3 = ParagraphStyle("H3", fontName="Helvetica-Bold", fontSize=10,
                         textColor=BLUE, spaceAfter=3, spaceBefore=10, leading=14)
    BODY = ParagraphStyle("Body", fontName="Helvetica", fontSize=9.5,
                           textColor=DARK_GRAY, spaceAfter=4, leading=14,
                           alignment=TA_JUSTIFY)
    BULLET = ParagraphStyle("Bullet", fontName="Helvetica", fontSize=9.5,
                              textColor=DARK_GRAY, spaceAfter=2, leading=13,
                              leftIndent=18, bulletIndent=6)
    CHECK = ParagraphStyle("Check", fontName="Helvetica", fontSize=9.5,
                            textColor=DARK_GRAY, spaceAfter=2, leading=13,
                            leftIndent=18)
    META = ParagraphStyle("Meta", fontName="Helvetica", fontSize=8.5,
                           textColor=MID_GRAY, spaceAfter=2, leading=12)
    COVER_META = ParagraphStyle("CoverMeta", fontName="Helvetica", fontSize=9,
                                 textColor=MID_GRAY, spaceAfter=3, leading=13,
                                 alignment=TA_CENTER)

    def sev_style(sev: str) -> ParagraphStyle:
        c = SEV_COLORS.get(sev.upper(), MID_GRAY)
        return ParagraphStyle(f"Sev_{sev}", fontName="Helvetica-Bold", fontSize=8.5,
                               textColor=WHITE, backColor=c, borderPad=(2, 4, 2, 4),
                               leading=12, alignment=TA_CENTER)

    # ── Build element list ────────────────────────────────────────────────────
    els = []

    lines = memo_markdown.strip().splitlines()
    i = 0
    in_cover = True  # treat first section (before first ##) as cover block

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── H1 ────────────────────────────────────────────────────────────
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title = stripped[2:].strip()
            # Remove leading "DUE DILIGENCE MEMO — "
            title = re.sub(r"^DUE DILIGENCE MEMO\s*[—-]\s*", "", title, flags=re.I)
            els.append(Spacer(1, 0.15 * inch))
            els.append(Paragraph("DUE DILIGENCE MEMO", H1))
            if title:
                els.append(Paragraph(title, H1))
            i += 1
            continue

        # ── Cover meta lines (**key:** value) ─────────────────────────────
        if in_cover and re.match(r"^\*\*.+\*\*", stripped):
            text = _inline_bold(stripped)
            els.append(Paragraph(text, COVER_META))
            i += 1
            continue

        # ── H2 section header ─────────────────────────────────────────────
        if stripped.startswith("## "):
            in_cover = False
            heading = stripped[3:].strip()
            els.append(Spacer(1, 0.08 * inch))
            els.append(Paragraph(f"  {heading}", H2))
            i += 1
            continue

        # ── H3 ────────────────────────────────────────────────────────────
        if stripped.startswith("### "):
            heading = stripped[4:].strip()
            els.append(Paragraph(_inline_bold(heading), H3))
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────
        if stripped == "---":
            els.append(Spacer(1, 0.04 * inch))
            els.append(HRFlowable(width="100%", thickness=0.4,
                                   color=RULE_GRAY, spaceAfter=4))
            i += 1
            continue

        # ── Markdown table ────────────────────────────────────────────────
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_table(table_lines)
            if rows:
                els.append(_make_table(rows, TH_BG, ALT_ROW, WHITE, DARK_GRAY, RULE_GRAY))
                els.append(Spacer(1, 0.08 * inch))
            continue

        # ── Bullet list item ──────────────────────────────────────────────
        if stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            checked = stripped.startswith("- [x]")
            text = stripped[6:].strip()
            mark = "☑" if checked else "☐"
            els.append(Paragraph(f"{mark}  {_inline_bold(text)}", CHECK))
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            els.append(Paragraph(f"•  {_inline_bold(text)}", BULLET))
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────
        if re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            num = re.match(r"^(\d+)\.", stripped).group(1)
            els.append(Paragraph(f"<b>{num}.</b>  {_inline_bold(text)}", BULLET))
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────
        if not stripped:
            els.append(Spacer(1, 0.05 * inch))
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────────
        text = _inline_bold(stripped)
        # italics
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        style = META if stripped.startswith("*(") else BODY
        els.append(Paragraph(text, style))
        i += 1

    doc.build(els, onFirstPage=_on_page, onLaterPages=_on_page)
    return buf.getvalue()


def _make_table(rows, th_bg, alt_bg, th_fg, body_fg, rule_color):
    """Build a styled ReportLab Table from a 2-D list of strings."""
    from reportlab.lib.units import inch
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    cell_style = ParagraphStyle("Cell", fontName="Helvetica", fontSize=8.5,
                                 textColor=body_fg, leading=12)
    head_style = ParagraphStyle("Head", fontName="Helvetica-Bold", fontSize=8.5,
                                 textColor=th_fg, leading=12)

    table_data = []
    for r_idx, row in enumerate(rows):
        s = head_style if r_idx == 0 else cell_style
        table_data.append([Paragraph(_inline_bold(c), s) for c in row])

    n_cols = max(len(r) for r in table_data)
    col_w = (7.0 * inch) / n_cols

    t = Table(table_data, colWidths=[col_w] * n_cols, repeatRows=1)

    style_cmds = [
        ("BACKGROUND",  (0, 0), (-1, 0), th_bg),
        ("TEXTCOLOR",   (0, 0), (-1, 0), th_fg),
        ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, -1), 8.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [alt_bg, "white"]),
        ("GRID",        (0, 0), (-1, -1), 0.4, rule_color),
        ("LEFTPADDING",  (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING",   (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 5),
        ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
    ]
    t.setStyle(TableStyle(style_cmds))
    return t


# ── DOCX via python-docx ──────────────────────────────────────────────────────

def to_docx(memo_markdown: str, summary: str = "", firm_name: str = "") -> bytes:
    """Convert a markdown IC memo to a styled Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Helper: set paragraph color + bold ───────────────────────────────────
    def _add_heading(text: str, level: int):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            if level == 1:
                run.font.color.rgb = RGBColor(0x0a, 0x27, 0x44)
                run.font.size = Pt(18)
            elif level == 2:
                run.font.color.rgb = RGBColor(0x1a, 0x5f, 0xa8)
                run.font.size = Pt(13)
        return p

    def _add_para(text: str, bold: bool = False, italic: bool = False,
                  size: int = 10, color: tuple = None, center: bool = False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Handle inline bold: **text**
        parts = re.split(r"(\*\*.+?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
                run.bold = bold
                run.italic = italic
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p

    def _add_bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        parts = re.split(r"(\*\*.+?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(10)
        return p

    def _add_table(rows: list[list[str]]):
        if not rows:
            return
        n_cols = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=n_cols)
        t.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            tr = t.rows[r_idx]
            for c_idx, cell_text in enumerate(row):
                if c_idx >= len(tr.cells):
                    break
                cell = tr.cells[c_idx]
                cell.text = _strip_md(cell_text)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if r_idx == 0:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                if r_idx == 0:
                    # shade header row navy
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "1a3a5c")
                    tcPr.append(shd)
        doc.add_paragraph()

    # ── Parse and render ──────────────────────────────────────────────────────
    lines = memo_markdown.strip().splitlines()
    i = 0

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if stripped.startswith("# ") and not stripped.startswith("## "):
            title = stripped[2:].strip()
            title = re.sub(r"^DUE DILIGENCE MEMO\s*[—-]\s*", "", title, flags=re.I)
            _add_heading("DUE DILIGENCE MEMO", 1)
            if title:
                _add_para(title, bold=True, size=14, center=True,
                          color=(0x1a, 0x5f, 0xa8))
            i += 1
            continue

        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            _add_heading(heading, 2)
            i += 1
            continue

        if stripped.startswith("### "):
            heading = stripped[4:].strip()
            _add_heading(heading, 3)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(_parse_table(table_lines))
            continue

        if stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            text = stripped[6:].strip()
            _add_bullet(f"☐  {text}")
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            _add_bullet(stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.add_run(text).font.size = Pt(10)
            i += 1
            continue

        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # meta / italic lines
        color = (0x55, 0x55, 0x55) if stripped.startswith("*(") else None
        _add_para(stripped, color=color)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
